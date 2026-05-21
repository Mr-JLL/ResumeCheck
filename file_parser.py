"""
file_parser.py
简历文件文字提取模块
支持：文字版PDF（PyMuPDF/fitz）、Word文档（python-docx）、扫描版PDF（PaddleOCR）
扫描版判断：有效汉字 < 50 时自动触发 OCR
Windows注意：pdf2image 需要 Poppler，会自动从常见路径查找或跳过
"""

import os
import re
import logging

logger = logging.getLogger(__name__)

# ── 可选依赖懒加载 ──────────────────────────────────────────
# 使用 PyMuPDF (fitz) 替代 pdfplumber，无 pdfminer.six 版本冲突
try:
    import fitz  # PyMuPDF
    HAS_PDFPLUMBER = True   # 变量名保持兼容，实际用 fitz
except ImportError:
    HAS_PDFPLUMBER = False
    logger.warning("PyMuPDF 未安装，文字版PDF解析不可用。pip install pymupdf")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx 未安装，Word解析不可用。pip install python-docx")

# PaddleOCR 懒加载（首次使用时初始化，避免启动耗时）
_ocr_engine = None
HAS_PADDLE = False
try:
    import paddleocr  # noqa: F401
    HAS_PADDLE = True
except ImportError:
    logger.warning("paddleocr 未安装，扫描版PDF将无法识别。pip install paddleocr paddlepaddle")

# pdf2image 懒加载
HAS_PDF2IMAGE = False
_POPPLER_PATH = None

def _init_pdf2image():
    """检测 pdf2image 和 Poppler 是否可用"""
    global HAS_PDF2IMAGE, _POPPLER_PATH
    try:
        from pdf2image import convert_from_path  # noqa: F401
        # 尝试常见的 Windows Poppler 安装路径（含嵌套解压结构）
        candidate_paths = [
            r"C:\poppler\bin\Release-26.02.0-0\poppler-26.02.0\Library\bin",
            r"C:\poppler\bin",
            r"C:\Program Files\poppler\bin",
            r"C:\Program Files (x86)\poppler\bin",
            os.path.join(os.path.dirname(os.path.abspath(__file__)), "poppler", "bin"),
        ]
        # 自动发现 C:\poppler\bin 下的任意嵌套版本目录
        base = r"C:\poppler\bin"
        if os.path.isdir(base):
            for entry in os.listdir(base):
                guessed = os.path.join(base, entry, "Library", "bin")
                if os.path.isdir(guessed):
                    candidate_paths.insert(0, guessed)
                nested = os.path.join(base, entry)
                for sub in os.listdir(nested) if os.path.isdir(nested) else []:
                    guessed2 = os.path.join(nested, sub, "Library", "bin")
                    if os.path.isdir(guessed2):
                        candidate_paths.insert(0, guessed2)
        import shutil
        if shutil.which("pdftoppm"):
            _POPPLER_PATH = None  # 已在 PATH 中
            HAS_PDF2IMAGE = True
        else:
            for p in candidate_paths:
                if os.path.exists(os.path.join(p, "pdftoppm.exe")):
                    _POPPLER_PATH = p
                    HAS_PDF2IMAGE = True
                    break
        if not HAS_PDF2IMAGE:
            logger.warning(
                "pdf2image 已安装但找不到 Poppler。"
                "请下载 Poppler for Windows 并放到 C:\\poppler\\bin 目录，"
                "或将 pdftoppm.exe 所在目录添加到 PATH。扫描版PDF将无法识别。"
            )
    except ImportError:
        logger.warning("pdf2image 未安装。pip install pdf2image")


_init_pdf2image()


def _get_ocr():
    """懒加载 PaddleOCR（首次调用时初始化，约10-30秒）"""
    global _ocr_engine
    if _ocr_engine is None:
        if not HAS_PADDLE:
            raise RuntimeError("paddleocr 未安装，无法处理扫描版PDF")
        from paddleocr import PaddleOCR
        logger.info("正在初始化 PaddleOCR 中文模型（首次约需30秒）...")
        _ocr_engine = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        logger.info("PaddleOCR 初始化完成")
    return _ocr_engine


def _is_pdf_scanned(text: str) -> bool:
    """有效汉字字符数 < 50 则判定为扫描版"""
    chinese_chars = re.findall(r'[一-鿿]', text)
    return len(chinese_chars) < 50


def _extract_pdf_text(file_path: str) -> str:
    """PyMuPDF (fitz) 提取文字版PDF"""
    if not HAS_PDFPLUMBER:
        raise RuntimeError("PyMuPDF 未安装，无法提取文字版PDF")
    text_parts = []
    with fitz.open(file_path) as pdf:
        for page in pdf:
            t = page.get_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def _extract_pdf_ocr(file_path: str) -> str:
    """pdf2image + PaddleOCR 提取扫描版PDF"""
    if not HAS_PDF2IMAGE:
        raise RuntimeError(
            "扫描版PDF识别需要 pdf2image 和 Poppler，请参考日志中的安装说明"
        )
    from pdf2image import convert_from_path
    import numpy as np

    kwargs = {"dpi": 200}
    if _POPPLER_PATH:
        kwargs["poppler_path"] = _POPPLER_PATH

    images = convert_from_path(file_path, **kwargs)
    ocr = _get_ocr()
    texts = []
    for img in images:
        img_array = np.array(img)
        result = ocr.ocr(img_array, cls=True)
        if result:
            for line in result:
                if line:
                    for word_info in line:
                        if word_info and len(word_info) >= 2:
                            texts.append(word_info[1][0])
    return "\n".join(texts)


def _extract_word_text(file_path: str) -> str:
    """python-docx 提取 Word 文档"""
    if not HAS_DOCX:
        raise RuntimeError("python-docx 未安装，无法提取Word文档")
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 同时提取表格中的文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_text_from_file(file_path: str) -> str:
    """
    自动检测文件类型并提取文字，返回提取的原始文字字符串。
    file_path：上传文件的本地路径（.pdf 或 .docx/.doc）
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx" or ext == ".doc":
        return _extract_word_text(file_path)

    if ext == ".pdf":
        # 先尝试文字版提取
        try:
            text = _extract_pdf_text(file_path)
        except Exception as e:
            logger.warning(f"文字版PDF提取失败 ({e})，尝试OCR")
            text = ""

        if _is_pdf_scanned(text):
            logger.info(f"检测到扫描版PDF（有效汉字<50），启动OCR: {os.path.basename(file_path)}")
            try:
                return _extract_pdf_ocr(file_path)
            except Exception as e:
                logger.warning(f"OCR失败 ({e})，返回部分文字版内容")
                return text or f"[OCR失败：{e}]"
        return text

    raise ValueError(f"不支持的文件格式：{ext}，仅支持 .pdf .docx .doc")
